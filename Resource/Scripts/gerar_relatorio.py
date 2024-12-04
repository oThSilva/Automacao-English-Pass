from docx import Document
from docx.shared import Inches

evidencias = {
    "English Pass/Results/Photos evidence/Evidência1.png": (
        "Cadastrar usuário com nome inválido",
        "Validar se ao cadastrar um usuário com nome inválido é apresentada uma mensagem amigável",
        "Passou",
    ),
    "English Pass/Results/Photos evidence/Evidência2.png": (
        "Cadastrar usuário com sobrenome inválido",
        "Validar se ao cadastrar um usuário com sobrenome inválido é apresentada uma mensagem amigável",
        "Passou",
    ),
    "English Pass/Results/Photos evidence/Evidência3.png": (
        "Cadastrar usuário com data de nascimento inválida",
        "Validar se ao cadastrar um usuário com data de nascimento inválida é apresentada uma mensagem amigável",
        "Passou",
    ),
    "English Pass/Results/Photos evidence/Evidência4.png": (
        "Cadastrar usuário com CPF inválido",
        "Validar se ao cadastrar um usuário com CPF inválido é apresentada uma mensagem amigável",
        "Passou",
    ),
    "English Pass/Results/Photos evidence/Evidência5.png": (
        "Cadastrar usuário com email inválido",
        "Validar se ao cadastrar um usuário com email inválido é apresentada uma mensagem amigável",
        "Passou",
    ),
    "English Pass/Results/Photos evidence/Evidência6.png": (
        "Cadastrar usuário com emails diferentes",
        "Validar se ao cadastrar um usuário com emails diferentes é apresentada uma mensagem amigável",
        "Passou",
    ),
    "English Pass/Results/Photos evidence/Evidência7.png": (
        "Cadastrar usuário com senha inválida",
        "Validar se ao cadastrar um usuário com senha inválida é apresentada uma mensagem amigável",
        "Passou",
    ),
    "English Pass/Results/Photos evidence/Evidência8.png": (
        "Cadastrar usuário com senhas diferentes",
        "Validar se ao cadastrar um usuário com senhas diferentes é apresentada uma mensagem amigável",
        "Passou",
    ),
    "English Pass/Results/Photos evidence/Evidência9.png": (
        "Cadastrar usuário sem aceitar os termos",
        "Validar se ao cadastrar um usuário sem aceitar os termos é apresentada uma mensagem amigável",
        "Passou",
    ),
    "English Pass/Results/Photos evidence/Evidência10.png": (
        "Cadastrar usuário com CEP inválido",
        "Validar se ao cadastrar um usuário com CEP inválido é apresentada uma mensagem amigável",
        "Passou",
    ),
    "English Pass/Results/Photos evidence/Evidência11.png": (
        "Cadastrar usuário com CEP que não exista",
        "Validar se ao cadastrar um usuário com CEP que não exista é apresentada uma mensagem amigável",
        "Passou",
    ),
    "English Pass/Results/Photos evidence/Evidência12.png": (
        "Cadastrar usuário com CPF já cadastrado",
        "Validar se ao cadastrar um usuário com CPF já cadastrado é apresentada uma mensagem amigável",
        "Passou",
    ),
    "English Pass/Results/Photos evidence/Evidência13.png": (
        "Cadastrar usuário com email já cadastrado",
        "Validar se ao cadastrar um usuário com email já cadastrado é apresentada uma mensagem amigável",
        "Passou",
    ),
    "English Pass/Results/Photos evidence/Evidência14.png": (
        "Cadastrar usuário no site English Pass",
        "Validar se é possível cadastrar um usuário no site English Pass",
        "Passou",
    ),
    "English Pass/Results/Photos evidence/Evidência15.png": (
        "Realizando o login e logout",
        "Validar se é possível realizar login e logout com usuário já cadastrado",
        "Passou",
    ),
}


def gerar_relatorio(evidencias, caminho_saida):
    doc = Document()
    doc.add_heading("Evidências dos testes automatizados", level=1)

    for index, (caminho_imagem, (nome, descricao, status)) in enumerate(
        evidencias.items(), start=1
    ):
        doc.add_heading(f"Evidência do cenário {index}. {nome}", level=2)
        doc.add_picture(caminho_imagem, width=Inches(7))
        doc.add_paragraph(f"Descrição do teste: {descricao}")
        doc.add_paragraph(f"Resultado do teste: {status}")

    doc.save(caminho_saida)
    print("Relatório gerado com sucesso!")


caminho_saida = "English Pass/Evidências dos testes.docx"
gerar_relatorio(evidencias, caminho_saida)
